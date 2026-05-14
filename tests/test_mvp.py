from pathlib import Path

from docx import Document

from draft_analyzer.cli import build_document
from draft_analyzer.metrics_rhetoric import CONTRIBUTION_MARKERS, has_marker
from draft_analyzer.metrics_syntax import detect_nominalizations
from draft_analyzer.sectioning import classify_section_heading


def build_fixture(path: Path) -> None:
    d = Document()
    d.add_heading("Abstract", level=1)
    d.add_paragraph("We contribute by identifying mechanisms in case study research using interviews.")
    d.add_heading("Introduction", level=1)
    d.add_paragraph("However, little is known about routine dynamics.")
    d.add_heading("Methods", level=1)
    d.add_paragraph("Data were collected through interviews and documents. We coded themes inductively.")
    d.add_heading("Findings", level=1)
    d.add_paragraph('"This changed everything," Participant 4 said.')
    d.add_paragraph("This indicates a coordination mechanism.")
    d.add_heading("Discussion", level=1)
    d.add_paragraph("We contribute to theory and discuss boundary conditions.")
    d.add_heading("References", level=1)
    d.add_paragraph("Author, A. (2020). Title.")
    d.save(path)


def test_build_document(tmp_path: Path):
    p = tmp_path / "draft.docx"
    build_fixture(p)
    doc = build_document(str(p))
    assert doc.sections
    assert any(s.section_type == "methods" for s in doc.sections)
    assert doc.global_metrics


def test_section_classification():
    assert classify_section_heading("Data and methods") == "methods"
    assert classify_section_heading("Bibliography") == "references"


def test_nominalization_detection():
    tokens = [("coordination", "NOUN"), ("run", "VERB")]
    out = detect_nominalizations(tokens, ("tion", "ment"))
    assert "coordination" in out


def test_contribution_marker():
    assert has_marker("We contribute to theory", CONTRIBUTION_MARKERS)
